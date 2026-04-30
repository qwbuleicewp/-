# ============================================================
# 模块：查重预览 (duplicate_preview.py) —— 修复滚动、允许文本选择
# ============================================================
import os
import tkinter as tk
from PIL import Image, ImageTk
import difflib
import utils
import config
from video_player import VideoPlayer

DIFFLIB_AVAILABLE = True

class DuplicatePreviewMixin:

    def _video_sync_callback(self, src_id, action, value):
        dst_id = 1 - src_id
        dst_player = self.video_players[dst_id]
        if not dst_player:
            return
        if action == 'play':
            dst_player.sync_play_state(value)
        elif action == 'seek':
            dst_player.sync_seek(value)

    def _update_preview(self):
        for i in range(2):
            if self.video_players[i]:
                self.video_players[i].destroy()
                self.video_players[i] = None
            self.archive_internal_state[i] = None
            label = self.preview_labels[i]
            label.pack_forget()
            label.config(image="", text="")
            if hasattr(label, 'original_image'):
                del label.original_image
            info = self.preview_infos[i]
            info.pack_forget()
            info.config(state='normal')
            info.delete(1.0, tk.END)
            self._make_text_readonly(info)

        paths = self.selected_preview_paths
        if paths[0] and paths[1] and self._is_text_file(paths[0]) and self._is_text_file(paths[1]):
            self._display_text_diff(paths[0], paths[1])
            return

        for i, path in enumerate(paths):
            self._display_file_in_preview(path, i)

    def _make_text_readonly(self, text_widget):
        """允许选择文本，但禁止修改内容"""
        text_widget.config(state='normal')
        text_widget.bind("<Key>", lambda e: "break")

    def _is_text_file(self, path):
        low = path.lower()
        return low.endswith(('.txt', '.md', '.py', '.json', '.xml', '.csv',
                             '.log', '.ini', '.cfg', '.html', '.css', '.js'))

    def _display_text_diff(self, path_a, path_b):
        try:
            with open(path_a, 'r', encoding='utf-8', errors='replace') as f:
                lines_a = f.read().splitlines()
        except:
            lines_a = ["[无法读取文件]"]
        try:
            with open(path_b, 'r', encoding='utf-8', errors='replace') as f:
                lines_b = f.read().splitlines()
        except:
            lines_b = ["[无法读取文件]"]

        if not DIFFLIB_AVAILABLE:
            for i, lines in enumerate([lines_a, lines_b]):
                info = self.preview_infos[i]
                info.config(state='normal')
                info.delete(1.0, tk.END)
                info.insert(1.0, '\n'.join(lines[:2000]))
                self._make_text_readonly(info)
                info.pack(fill=tk.BOTH, expand=True)
                self._bind_sync_scroll(info, i)
            return

        sm = difflib.SequenceMatcher(None, lines_a, lines_b)
        aligned_a, aligned_b = [], []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                for line in lines_a[i1:i2]:
                    aligned_a.append(('equal', line))
                    aligned_b.append(('equal', line))
            elif tag == 'replace':
                for line in lines_a[i1:i2]:
                    aligned_a.append(('replace', line))
                for line in lines_b[j1:j2]:
                    aligned_b.append(('replace', line))
                len_a, len_b = i2-i1, j2-j1
                if len_a > len_b:
                    for _ in range(len_a - len_b):
                        aligned_b.append(('empty', ''))
                elif len_b > len_a:
                    for _ in range(len_b - len_a):
                        aligned_a.append(('empty', ''))
            elif tag == 'delete':
                for line in lines_a[i1:i2]:
                    aligned_a.append(('delete', line))
                    aligned_b.append(('empty', ''))
            elif tag == 'insert':
                for line in lines_b[j1:j2]:
                    aligned_a.append(('empty', ''))
                    aligned_b.append(('insert', line))

        for info in self.preview_infos:
            info.tag_configure("equal", background="white", foreground="black")
            info.tag_configure("replace", background="#ffcccc", foreground="black")
            info.tag_configure("delete", background="white", foreground="blue")
            info.tag_configure("insert", background="white", foreground="blue")
            info.tag_configure("empty", background="white", foreground="black")

        for i, aligned in enumerate([aligned_a, aligned_b]):
            info = self.preview_infos[i]
            info.config(state='normal')
            info.delete(1.0, tk.END)
            for tag, line in aligned[:2000]:
                if tag == 'empty':
                    info.insert(tk.END, '\n', tag)
                else:
                    info.insert(tk.END, line + '\n', tag)
            self._make_text_readonly(info)
            info.pack(fill=tk.BOTH, expand=True)
            self._bind_sync_scroll(info, i)

        self.preview_title_a.config(text=f"预览 A: {os.path.basename(path_a)}")
        self.preview_title_b.config(text=f"预览 B: {os.path.basename(path_b)}")

    def _bind_sync_scroll(self, text_widget, idx):
        # 绑定滚动事件
        text_widget.bind("<MouseWheel>", lambda e: self._on_sync_scroll(e, idx))
        text_widget.bind("<Button-4>", lambda e: self._on_sync_scroll(e, idx))
        text_widget.bind("<Button-5>", lambda e: self._on_sync_scroll(e, idx))

    def _on_sync_scroll(self, event, src_idx):
        """同步滚动：两个预览区滚动相同单位"""
        if self._syncing_scroll:
            return
        self._syncing_scroll = True
        try:
            # 计算滚动单位
            if event.delta:
                # Windows: delta 为 120 的倍数，向上滚动为正
                delta = -1 if event.delta > 0 else 1
            elif event.num == 4:   # Linux 向上
                delta = -1
            elif event.num == 5:   # Linux 向下
                delta = 1
            else:
                delta = 0

            # 两个预览区同时滚动
            for i in (0, 1):
                info = self.preview_infos[i]
                if info.winfo_exists():
                    info.yview_scroll(delta, "units")

            # 阻止系统默认滚动行为
            return "break"
        finally:
            self._syncing_scroll = False

    def _display_file_in_preview(self, path, preview_index, title_extra=""):
        i = preview_index
        title = self.preview_titles[i]
        label = self.preview_labels[i]
        info = self.preview_infos[i]

        label.pack_forget()
        info.pack_forget()
        label.config(image="", text="")
        if hasattr(label, 'original_image'):
            del label.original_image
        if self.video_players[i]:
            self.video_players[i].destroy()
            self.video_players[i] = None

        if not path or not os.path.exists(path):
            title.config(text=f"预览 {chr(65+i)}")
            info.pack(fill=tk.BOTH, expand=True)
            info.config(state='normal')
            info.delete(1.0, tk.END)
            self._make_text_readonly(info)
            return

        title_text = path if not title_extra else f"{path} {title_extra}"
        title.config(text=title_text)
        low = path.lower()

        # --- 图片 ---
        if low.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
            label.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            try:
                img = Image.open(path)
                label.update_idletasks()
                max_w = label.winfo_width() or 400
                max_h = label.winfo_height() or 400
                img.thumbnail((max_w, max_h))
                photo = ImageTk.PhotoImage(img)
                label.config(image=photo)
                label.image = photo
                label.original_image = img
                label.scale_factor = 1.0
                label.bind("<Control-MouseWheel>",
                           lambda e, lbl=label, idx=i: self._zoom_image_sync(e, lbl, idx))
                label.bind("<Double-Button-1>", lambda e, p=path: utils.open_file_or_folder(p))
            except Exception as e:
                label.config(image="", text="预览失败")

        # --- 视频 ---
        elif low.endswith(('.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv', '.webm')):
            label.pack(fill=tk.BOTH, expand=True)
            player = VideoPlayer(label.master, path, 400, 250,
                                 sync_callback=self._video_sync_callback, player_id=i)
            player.pack(fill=tk.BOTH, expand=True)
            self.video_players[i] = player
            player.video_label.bind("<Double-Button-1>", lambda e, p=path: utils.open_file_or_folder(p))

        # --- 其他文本文件 ---
        else:
            info.pack(fill=tk.BOTH, expand=True)
            content = self._get_file_preview_content(path)
            info.config(state='normal')
            info.delete(1.0, tk.END)
            info.insert(1.0, content)
            self._make_text_readonly(info)
            self._bind_sync_scroll(info, i)
            if low.endswith('.pdf'):
                info.bind("<Double-Button-1>", lambda e, p=path: self._open_all_pdfs_in_group(p))
            else:
                info.bind("<Double-Button-1>", lambda e, p=path: utils.open_file_or_folder(p))

    # ---------- 图片拖拽与缩放（保持不变） ----------
    def _zoom_image(self, event, label):
        if not hasattr(label, 'original_image'):
            return
        delta = event.delta
        if delta > 0:
            label.scale_factor *= 1.1
        else:
            label.scale_factor *= 0.9
        label.scale_factor = max(0.1, min(5.0, label.scale_factor))
        img = label.original_image
        new_w = int(img.width * label.scale_factor)
        new_h = int(img.height * label.scale_factor)
        resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        photo = ImageTk.PhotoImage(resized)
        label.config(image=photo)
        label.image = photo

    def _zoom_image_sync(self, event, label, src_idx):
        self._zoom_image(event, label)
        dst_idx = 1 - src_idx
        dst_label = self.preview_labels[dst_idx]
        if hasattr(dst_label, 'original_image') and dst_label.winfo_ismapped():
            dst_label.scale_factor = label.scale_factor
            img = dst_label.original_image
            new_w = int(img.width * dst_label.scale_factor)
            new_h = int(img.height * dst_label.scale_factor)
            resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(resized)
            dst_label.config(image=photo)
            dst_label.image = photo

    def _sync_video_progress(self, src_idx):
        src_player = self.video_players[src_idx]
        dst_player = self.video_players[1 - src_idx]
        if src_player and dst_player:
            dst_player.progress_var.set(src_player.progress_var.get())
            dst_player._show_frame(int(src_player.current_frame))

    def _open_all_pdfs_in_group(self, pdf_path):
        group = self._find_group_by_path(pdf_path)
        if not group:
            return
        pdf_files = [f['path'] for f in group if f['path'].lower().endswith('.pdf')]
        for p in pdf_files:
            utils.open_file_or_folder(p)

    def _get_file_preview_content(self, path):
        low = path.lower()
        if low.endswith('.pdf'):
            return self._preview_pdf_content(path)
        elif low.endswith(('.zip', '.rar', '.7z', '.tar', '.gz', '.bz2')):
            return self._preview_archive_content(path)
        elif low.endswith(('.txt', '.md', '.py', '.json', '.xml', '.csv', '.log',
                           '.ini', '.cfg', '.html', '.css', '.js')):
            return self._preview_text_file(path)
        elif low.endswith(('.docx', '.doc')):
            return self._preview_word_file(path)
        elif low.endswith(('.xlsx', '.xls')):
            return self._preview_excel_file(path)
        else:
            return "双击预览区用默认程序打开"

    def _preview_text_file(self, path, max_chars=2000):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read(max_chars)
            return content + ("..." if len(content) == max_chars else "")
        except Exception as e:
            return f"无法读取文件: {e}"

    def _preview_pdf_content(self, path, max_chars=2000):
        try:
            import PyPDF2
            text = ""
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for i in range(min(5, len(reader.pages))):
                    page = reader.pages[i].extract_text()
                    if page:
                        text += page + "\n"
            if not text:
                return "（PDF 无文本内容）"
            return text[:max_chars] + ("..." if len(text) > max_chars else "")
        except ImportError:
            return "未安装 PyPDF2 库"
        except Exception as e:
            return f"PDF 解析失败: {e}"

    def _preview_word_file(self, path):
        try:
            import docx
            doc = docx.Document(path)
            text = "\n".join(p.text for p in doc.paragraphs[:20])
            if not text:
                for table in doc.tables:
                    for row in table.rows:
                        text += " | ".join(cell.text for cell in row.cells) + "\n"
            return text if text else "（文档无文本内容）"
        except ImportError:
            return "未安装 python-docx 库"
        except Exception as e:
            return f"Word 解析失败: {e}"

    def _preview_excel_file(self, path):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            sheet = wb.active
            rows = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= 50:
                    rows.append("...")
                    break
                rows.append(" | ".join(str(cell) if cell is not None else "" for cell in row))
            wb.close()
            return "\n".join(rows) if rows else "（表格无数据）"
        except ImportError:
            return "未安装 openpyxl 库"
        except Exception as e:
            return f"Excel 解析失败: {e}"

    def _preview_archive_content(self, path):
        if not hasattr(self, '_archive_cache'):
            self._archive_cache = {}
        ext = os.path.splitext(path)[1].lower()
        lines = []
        try:
            if ext == '.zip':
                import zipfile
                with zipfile.ZipFile(path, 'r') as zf:
                    for name in zf.namelist()[:100]:
                        info = zf.getinfo(name)
                        lines.append(f"  {name} ({utils.format_size(info.file_size)})")
            elif ext in ('.tar', '.gz', '.bz2'):
                import tarfile
                mode = 'r:gz' if ext == '.gz' else 'r:bz2' if ext == '.bz2' else 'r'
                with tarfile.open(path, mode) as tf:
                    for member in tf.getmembers()[:100]:
                        if member.isfile():
                            lines.append(f"  {member.name} ({utils.format_size(member.size)})")
            elif ext == '.rar':
                try:
                    import rarfile
                    with rarfile.RarFile(path, 'r') as rf:
                        for name in rf.namelist()[:100]:
                            info = rf.getinfo(name)
                            lines.append(f"  {name} ({utils.format_size(info.file_size)})")
                except ImportError:
                    return "未安装 rarfile 库"
            elif ext == '.7z':
                try:
                    import py7zr
                    with py7zr.SevenZipFile(path, 'r') as szf:
                        for name, info in szf.getnames()[:100]:
                            lines.append(f"  {name} ({utils.format_size(info.size)})")
                except ImportError:
                    return "未安装 py7zr 库"
            else:
                return "不支持的压缩格式"
            if not lines:
                return "（压缩包内无文件）"
            return "\n".join(lines)
        except Exception as e:
            if 'password' in str(e).lower() or 'encrypted' in str(e).lower():
                return "压缩包已加密"
            return f"读取失败: {e}"
